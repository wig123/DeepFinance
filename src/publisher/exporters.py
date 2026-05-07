"""Publisher 导出器"""

import logging
import re
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Optional

import markdown
from jinja2 import Template

from .schemas import ExportFormat, ReportMetadata
from .templates import HTML_TEMPLATE, PDF_CSS

logger = logging.getLogger(__name__)


class BaseExporter(ABC):
    """导出器基类"""

    format: ExportFormat

    @abstractmethod
    def export(
        self,
        content: str,
        output_path: Path,
        metadata: ReportMetadata,
        sources: Optional[list[str]] = None,
    ) -> bool:
        """导出报告"""
        pass


class HTMLExporter(BaseExporter):
    """HTML 导出器"""

    format = ExportFormat.HTML

    def __init__(self):
        self.md = markdown.Markdown(
            extensions=[
                "tables",
                "fenced_code",
                "footnotes",
                "toc",
                "attr_list",
            ]
        )

    def export(
        self,
        content: str,
        output_path: Path,
        metadata: ReportMetadata,
        sources: Optional[list[str]] = None,
    ) -> bool:
        """导出为 HTML"""
        try:
            # 转换 Markdown 为 HTML
            html_content = self.md.convert(content)
            self.md.reset()

            # 渲染模板
            template = Template(HTML_TEMPLATE)
            full_html = template.render(
                title=metadata.title,
                content=html_content,
                sources=sources or [],
                generated_at=datetime.now().strftime("%Y-%m-%d %H:%M"),
            )

            # 写入文件
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(full_html, encoding="utf-8")

            logger.info(f"[Publisher] HTML 导出成功: {output_path}")
            return True

        except Exception as e:
            logger.error(f"[Publisher] HTML 导出失败: {e}")
            return False


class PDFExporter(BaseExporter):
    """PDF 导出器（使用 WeasyPrint）"""

    format = ExportFormat.PDF

    def __init__(self):
        self.html_exporter = HTMLExporter()

    def export(
        self,
        content: str,
        output_path: Path,
        metadata: ReportMetadata,
        sources: Optional[list[str]] = None,
    ) -> bool:
        """导出为 PDF"""
        try:
            # 尝试导入 weasyprint
            try:
                from weasyprint import HTML, CSS
            except ImportError:
                logger.warning("[Publisher] WeasyPrint 未安装，跳过 PDF 导出")
                logger.info("安装命令: pip install weasyprint")
                return False

            # 先生成 HTML
            html_path = output_path.with_suffix(".html")
            if not self.html_exporter.export(content, html_path, metadata, sources):
                return False

            # 转换为 PDF
            html_doc = HTML(filename=str(html_path))
            pdf_css = CSS(string=PDF_CSS.replace("{{ title }}", metadata.title))

            html_doc.write_pdf(
                str(output_path),
                stylesheets=[pdf_css],
            )

            logger.info(f"[Publisher] PDF 导出成功: {output_path}")
            return True

        except Exception as e:
            logger.error(f"[Publisher] PDF 导出失败: {e}")
            return False


class DocxExporter(BaseExporter):
    """Word 文档导出器（使用 python-docx）"""

    format = ExportFormat.DOCX

    def export(
        self,
        content: str,
        output_path: Path,
        metadata: ReportMetadata,
        sources: Optional[list[str]] = None,
    ) -> bool:
        """导出为 Word 文档"""
        try:
            # 尝试导入 python-docx
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                logger.warning("[Publisher] python-docx 未安装，跳过 Word 导出")
                logger.info("安装命令: pip install python-docx")
                return False

            # 创建文档
            doc = Document()

            # 设置标题
            title = doc.add_heading(metadata.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加元信息
            doc.add_paragraph(f"作者: {metadata.author}")
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph()

            # 解析 Markdown 并添加内容
            self._parse_markdown_to_docx(doc, content)

            # 添加参考文献
            if sources:
                doc.add_page_break()
                doc.add_heading("参考文献", level=1)
                for source in sources:
                    p = doc.add_paragraph(source)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(-0.5)

            # 保存文档
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            logger.info(f"[Publisher] Word 导出成功: {output_path}")
            return True

        except Exception as e:
            logger.error(f"[Publisher] Word 导出失败: {e}")
            return False

    def _parse_markdown_to_docx(self, doc, content: str):
        """简单的 Markdown 到 Word 转换"""
        lines = content.split("\n")
        current_para = []

        for line in lines:
            # 标题
            if line.startswith("# "):
                self._flush_paragraph(doc, current_para)
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                self._flush_paragraph(doc, current_para)
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                self._flush_paragraph(doc, current_para)
                doc.add_heading(line[4:], level=3)
            # 列表
            elif line.startswith("- ") or line.startswith("* "):
                self._flush_paragraph(doc, current_para)
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                self._flush_paragraph(doc, current_para)
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            # 空行
            elif not line.strip():
                self._flush_paragraph(doc, current_para)
            # 普通段落
            else:
                current_para.append(line)

        self._flush_paragraph(doc, current_para)

    def _flush_paragraph(self, doc, para_lines: list):
        """输出段落"""
        if para_lines:
            text = " ".join(para_lines)
            # 移除 Markdown 格式标记
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # 粗体
            text = re.sub(r"\*(.*?)\*", r"\1", text)  # 斜体
            text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)  # 链接
            doc.add_paragraph(text)
            para_lines.clear()


class MarkdownExporter(BaseExporter):
    """Markdown 导出器（原样保存）"""

    format = ExportFormat.MARKDOWN

    def export(
        self,
        content: str,
        output_path: Path,
        metadata: ReportMetadata,
        sources: Optional[list[str]] = None,
    ) -> bool:
        """导出为 Markdown"""
        try:
            # 添加元信息头
            header = f"""---
title: {metadata.title}
author: {metadata.author}
date: {datetime.now().strftime('%Y-%m-%d')}
---

"""
            # 添加参考文献
            footer = ""
            if sources:
                footer = "\n\n---\n\n## 参考文献\n\n" + "\n".join(sources)

            full_content = header + content + footer

            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(full_content, encoding="utf-8")

            logger.info(f"[Publisher] Markdown 导出成功: {output_path}")
            return True

        except Exception as e:
            logger.error(f"[Publisher] Markdown 导出失败: {e}")
            return False


# 导出器注册表
EXPORTERS: dict[ExportFormat, type[BaseExporter]] = {
    ExportFormat.HTML: HTMLExporter,
    ExportFormat.PDF: PDFExporter,
    ExportFormat.DOCX: DocxExporter,
    ExportFormat.MARKDOWN: MarkdownExporter,
}


def get_exporter(format: ExportFormat) -> BaseExporter:
    """获取导出器实例"""
    exporter_class = EXPORTERS.get(format)
    if not exporter_class:
        raise ValueError(f"不支持的导出格式: {format}")
    return exporter_class()
